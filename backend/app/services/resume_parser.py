import io
import logging
from pathlib import Path
from typing import Optional, Tuple
import pdfplumber
import PyPDF2
from docx import Document
from app.core.config import settings

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc"}


class FileParsingError(Exception):
    pass


class FileValidationError(Exception):
    pass


def validate_file(file_bytes: bytes, filename: str) -> Tuple[bool, str, str]:
    """Validate uploaded resume file size and extension."""
    file_size_bytes = len(file_bytes)
    if file_size_bytes == 0:
        return False, "Uploaded file is empty. Please upload a valid resume.", ""

    if file_size_bytes > settings.MAX_FILE_SIZE_BYTES:
        size_mb = file_size_bytes / (1024 * 1024)
        return (
            False,
            f"File size ({size_mb:.2f} MB) exceeds maximum allowed {settings.MAX_FILE_SIZE_MB} MB.",
            "",
        )

    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        return (
            False,
            f"Unsupported file extension '{ext}'. Please upload a PDF (.pdf) or Word document (.docx).",
            "",
        )

    return True, "", ext.lstrip(".")


def _extract_pdf_hyperlinks(file_bytes: bytes) -> str:
    """Extract hyperlinks from PDF annotations (LinkedIn, GitHub, Portfolios)."""
    urls = []
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            if "/Annots" not in page:
                continue
            for annot_ref in page["/Annots"]:
                try:
                    annot = annot_ref.get_object()
                    if annot.get("/Subtype") != "/Link":
                        continue
                    action = annot.get("/A", {})
                    uri = action.get("/URI", "")
                    if uri:
                        if isinstance(uri, bytes):
                            uri = uri.decode("utf-8", errors="ignore")
                        uri = uri.strip()
                        if uri.startswith("http"):
                            urls.append(uri)
                except Exception:
                    pass
    except Exception:
        pass
    return "\n".join(list(set(urls)))


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF using pdfplumber with PyPDF2 fallback."""
    text = ""
    # Method 1: pdfplumber
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}. Trying PyPDF2...")

    # Method 2: PyPDF2 fallback if text is sparse or failed
    if not text.strip():
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                p_text = page.extract_text()
                if p_text:
                    text += p_text + "\n"
        except Exception as e:
            logger.error(f"PyPDF2 extraction failed: {e}")

    if not text.strip():
        raise FileParsingError(
            "Could not extract readable text from PDF. The document may be scanned, image-only, or encrypted."
        )

    # Append any extracted external hyperlinks
    links = _extract_pdf_hyperlinks(file_bytes)
    if links:
        text = text.strip() + "\n" + links

    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX including tables and hyperlinks."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []

        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        text = "\n".join(parts)

        # Extract hyperlinks
        try:
            for rel in doc.part.rels.values():
                if "hyperlink" in rel.reltype.lower():
                    url = rel._target
                    if isinstance(url, str) and url.startswith("http"):
                        text += "\n" + url
        except Exception:
            pass

        if not text.strip():
            raise FileParsingError("No text found in the DOCX file.")

        return text.strip()
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise FileParsingError(f"Failed to parse DOCX file: {e}") from e


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Unified text extraction dispatcher."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        raise FileValidationError(f"Unsupported file extension: {ext}")


def parse_resume_file(file_bytes: bytes, filename: str) -> Tuple[str, dict]:
    """Validate and extract text from an uploaded resume file."""
    is_valid, error_msg, file_type = validate_file(file_bytes, filename)
    if not is_valid:
        raise FileValidationError(error_msg)

    text = extract_text(file_bytes, filename)
    metadata = {
        "filename": filename,
        "file_type": file_type,
        "file_size_bytes": len(file_bytes),
        "text_length": len(text),
    }
    return text, metadata
